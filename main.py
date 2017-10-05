from docx import Document
from docx.shared import Pt
from pprint import pprint
import quickstart
import datetime
from datetime import datetime

# gets data from the sheet
data = quickstart.main()['values']

# custom variables for your document
title_of_template_document = 'INSERT TITLE OF TEMPLATE DOCUMENT'
row_identifier = 'INSERT IDENTIFIER IF YOU WANT TO USE ONLY CERTAIN ROWS'
search_text = '__________________'
count_of_items_to_replace = 6

# main runner
for entry in data:
    if entry[2] == row_identifier:
		document = Document(title_of_template_document)
		# gets the specific data you need
		replacement_text_array = [entry[0], entry[1], entry[6], entry[3], entry[7], entry[2]]
		for i in range(count_of_items_to_replace):
			old_text = document.paragraphs[i+1].text
			new_text = old_text.replace(search_text, replacement_text_array[i])
			document.paragraphs[i+1].text = new_text
		document.save('Document Title ' + replacement_text_array[0] + '.docx')